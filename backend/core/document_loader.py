import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.append(str(Path(__file__).parent.parent.parent))

import os
import logging
import re
from typing import List, Dict, Any
from llama_index.core import Document, VectorStoreIndex, ServiceContext
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.openai import OpenAIEmbedding
from supabase import create_client, Client
from backend.core.config import SUPABASE_URL, SUPABASE_KEY, OPENAI_API_KEY, EMBEDDING_MODEL
import time
import hashlib

logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')

def get_supabase_client(max_retries=3, retry_delay=1):
    """Initialize Supabase client with retry logic."""
    for attempt in range(max_retries):
        try:
            client = create_client(SUPABASE_URL, SUPABASE_KEY)
            # Test connection by trying to access a table
            client.table('ps-index').select('id').limit(1).execute()
            return client
        except Exception as e:
            if attempt < max_retries - 1:
                logging.warning(f"Failed to connect to Supabase (attempt {attempt + 1}/{max_retries}): {e}")
                time.sleep(retry_delay)
            else:
                logging.error(f"Failed to connect to Supabase after {max_retries} attempts: {e}")
                raise

# Initialize Supabase client with retry logic
supabase_client = get_supabase_client()

def ensure_table_exists(table_name: str):
    """Ensure Supabase table exists with correct configuration."""
    try:
        # Check if table exists by trying to query it
        try:
            supabase_client.table(table_name).select('id').limit(1).execute()
            logging.info(f"✓ Table {table_name} already exists")
            return True
        except Exception as e:
            # Table doesn't exist
            logging.error(f"✗ Table {table_name} does not exist")
            logging.error(f"Error: {e}")
            logging.info(f"ℹ Please create table {table_name} manually in Supabase dashboard:")
            logging.info(f"  1. Go to: Database > Tables > New Table")
            logging.info(f"  2. Name: {table_name}")
            logging.info(f"  3. Columns:")
            logging.info(f"     - id: bigint, primary key, auto-increment")
            logging.info(f"     - content: text")
            logging.info(f"     - metadata: jsonb")
            logging.info(f"     - embedding: vector(1536)")
            logging.info(f"  4. Enable Row Level Security: No")
            logging.info(f"  5. Click 'Save'")
            logging.info(f"  6. Enable pgvector extension in Database > Extensions")
            return False
            
    except Exception as e:
        logging.error(f"Error checking table existence: {e}")
        return False

def extract_metadata_from_filename(filename: str) -> Dict[str, Any]:
    """Extract metadata from filename."""
    section = extract_section_from_filename(filename)
    metadata = {
        "file_name": filename,
        "section": section
    }
    return metadata

def extract_section_from_filename(filename: str) -> str:
    """Extract section information from filename and content structure."""
    # Remove file extension
    name_without_ext = os.path.splitext(filename)[0]
    
    # Try to match section patterns like "1.2", "Section 1", "Appendix A", etc.
    patterns = [
        r"(\d+\.\d+)",  # 1.2, 2.3, etc.
        r"section\s*(\d+)",  # Section 1, section 2, etc.
        r"appendix\s*([A-Z])",  # Appendix A, Appendix B, etc.
        r"chapter\s*(\d+)",  # Chapter 1, Chapter 2, etc.
        r"part\s*(\d+)",  # Part 1, Part 2, etc.
    ]
    
    for pattern in patterns:
        match = re.search(pattern, name_without_ext, re.IGNORECASE)
        if match:
            return match.group(1)
    
    # If no pattern matches, try to extract meaningful parts from filename
    # Common clinical trial document patterns
    if "protocol" in name_without_ext.lower():
        if "summary" in name_without_ext.lower():
            return "Protocol Summary"
        elif "template" in name_without_ext.lower():
            return "Protocol Template"
        else:
            return "Protocol"
    elif "investigator" in name_without_ext.lower() and "brochure" in name_without_ext.lower():
        return "Investigator's Brochure"
    elif "reference" in name_without_ext.lower():
        return "Reference Protocol"
    elif "consent" in name_without_ext.lower():
        return "Informed Consent"
    elif "case" in name_without_ext.lower() and "report" in name_without_ext.lower():
        return "Case Report Form"
    
    # Default to filename without extension if no pattern matches
    return name_without_ext


def build_index(folder_path: str, collection_name: str) -> VectorStoreIndex:
    """
    Build a vector index from documents in the specified folder.
    
    Args:
        folder_path: Path to the folder containing documents
        collection_name: Name of the Supabase table to use
        
    Returns:
        VectorStoreIndex: The built index
    """
    try:
        logging.info(f"Indexing folder: {folder_path} into table: {collection_name}")
        
        # Check if folder exists
        if not os.path.exists(folder_path):
            logging.error(f"Folder {folder_path} does not exist")
            return None
            
        # Ensure table exists
        if not ensure_table_exists(collection_name):
            return None
            
        files = os.listdir(folder_path)
        logging.info(f"Found {len(files)} files in folder: {folder_path}")
        
        # Load and process documents
        docs = []
        for filename in files:
            path = os.path.join(folder_path, filename)
            if filename.endswith('.pdf'):
                logging.info(f"Processing PDF: {filename}")
                text = extract_text_from_pdf(path)
            elif filename.endswith('.docx'):
                logging.info(f"Processing DOCX: {filename}")
                text = extract_text_from_docx(path)
            else:
                logging.info(f"Skipping unsupported file: {filename}")
                continue
                
            logging.info(f"Extracted text length for {filename}: {len(text)}")
            if not text or not isinstance(text, str) or not text.strip():
                logging.warning(f"WARNING: Invalid text extracted for file: {filename}")
                continue
                
            # Use simple filename-based section detection (keeps existing behavior)
            section = extract_section_from_filename(filename)
            docs.append(Document(
                text=text, 
                metadata={
                    "section": section,
                    "filename": filename
                }
            ))
            
        logging.info(f"Loaded {len(docs)} documents from {folder_path}")
        
        if not docs:
            logging.error("No valid documents found to index")
            return None
            
        # Initialize embedding model
        embed_model = OpenAIEmbedding(model=EMBEDDING_MODEL, api_key=OPENAI_API_KEY)
        logging.info(f"Using embedding model: {EMBEDDING_MODEL}")
        
        # Chunk docs with proper overlap
        splitter = SentenceSplitter(chunk_size=1024, chunk_overlap=100)
        nodes = splitter.get_nodes_from_documents(docs)
        logging.info(f"Created {len(nodes)} chunks from documents")
        
        if not nodes:
            logging.error("No chunks created from documents")
            return None
            
        # Generate embeddings and upload to Supabase
        for i, node in enumerate(nodes):
            try:
                # Get node content and validate
                content = node.get_content()
                if not content or not isinstance(content, str) or not content.strip():
                    logging.warning(f"Skipping invalid chunk {i}")
                    continue

                # Compute hash for deduplication (content + filename)
                filename = node.metadata.get("filename", "") if node.metadata else ""
                hash_input = (content + filename).encode("utf-8")
                chunk_hash = hashlib.sha256(hash_input).hexdigest()

                # Check for existing chunk with same hash
                try:
                    existing = supabase_client.table(collection_name).select('id').eq('hash', chunk_hash).limit(1).execute()
                    if existing.data and len(existing.data) > 0:
                        logging.info(f"Duplicate chunk detected (hash: {chunk_hash}), skipping insert.")
                        continue
                except Exception as e:
                    logging.warning(f"Error checking for duplicate chunk: {e}")

                # Generate embedding
                try:
                    embedding = embed_model.get_text_embedding(content)
                    if not embedding:
                        logging.warning(f"Failed to generate embedding for chunk {i}")
                        continue
                except Exception as e:
                    logging.error(f"Error generating embedding for chunk {i}: {e}")
                    continue

                # Prepare data for Supabase
                data = {
                    "content": content,
                    "metadata": node.metadata or {},
                    "embedding": embedding,
                    "hash": chunk_hash
                }

                # Insert into Supabase
                try:
                    supabase_client.table(collection_name).insert(data).execute()
                except Exception as e:
                    logging.error(f"Error inserting chunk {i} into Supabase: {e}")
                    continue

                if (i + 1) % 100 == 0:
                    logging.info(f"Processed {i + 1} chunks")

            except Exception as e:
                logging.error(f"Error processing chunk {i}: {e}")
                continue
                
        logging.info(f"Uploaded chunks to table {collection_name}")
            
        # Create a simple index for querying - we'll use direct Supabase queries instead
        # of SupabaseVectorStore since the connection string approach is causing issues
        
        # Create a simple in-memory index for now
        # The actual vector search will be done through direct Supabase queries
        embed_model = OpenAIEmbedding(model=EMBEDDING_MODEL, api_key=OPENAI_API_KEY)
        
        # Create a dummy document to initialize the index
        dummy_doc = Document(text="dummy", metadata={})
        index = VectorStoreIndex.from_documents([dummy_doc], embed_model=embed_model)
        
        return index
        
    except Exception as e:
        logging.error(f"Error building index: {e}")
        return None

from PyPDF2 import PdfReader
import docx

def extract_text_from_pdf(filepath):
    reader = PdfReader(filepath)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_docx(filepath):
    doc = docx.Document(filepath)
    return "\n".join(para.text for para in doc.paragraphs)


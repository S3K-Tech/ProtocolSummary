from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import re

def sanitize_filename(filename: str) -> str:
    """Sanitize filename by removing/replacing invalid characters"""
    # Remove or replace invalid characters for filenames
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    # Remove extra spaces and limit length
    filename = re.sub(r'\s+', '_', filename.strip())
    return filename[:100]  # Limit to 100 characters

def create_word_doc(text: str, title: str) -> bytes:
    """Create a Word document with the given text and title."""
    try:
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Split text into paragraphs and add them
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():  # Only add non-empty paragraphs
                doc.add_paragraph(para.strip())
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
    except Exception as e:
        raise Exception(f"Failed to create Word document: {str(e)}")

def create_pdf(text: str, title: str) -> bytes:
    """Create a PDF with the given text and title."""
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        story = []
        
        # Add title
        story.append(Paragraph(title, styles['Title']))
        story.append(Paragraph("<br/><br/>", styles['Normal']))  # Add spacing
        
        # Split text into paragraphs and add them
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():  # Only add non-empty paragraphs
                # Clean up the text for PDF (handle special characters)
                clean_para = para.strip().replace('\n', '<br/>')
                story.append(Paragraph(clean_para, styles['Normal']))
                story.append(Paragraph("<br/>", styles['Normal']))  # Add spacing between paragraphs
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        raise Exception(f"Failed to create PDF document: {str(e)}")
